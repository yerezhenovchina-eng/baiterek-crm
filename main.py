# main.py — НИХ Байтерек corporate dashboard
import io, bcrypt, os
from datetime import datetime
from fastapi import Query, FastAPI, Request, Form, File, UploadFile, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from itsdangerous import URLSafeTimedSerializer, BadSignature
import uvicorn

from database import init_db, get_db, ORGANIZATIONS, ORG_MAP, ORG_BY_USERNAME, now

app = FastAPI(title="НИХ Байтерек — Корпоративная карточка")
templates = Jinja2Templates(directory="templates")

# Static files (logos etc.)
if os.path.isdir("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")
SECRET_KEY = "baiterek-secret-2024"
SESSION_MAX_AGE = 8 * 3600
serializer = URLSafeTimedSerializer(SECRET_KEY)

def fmt(d):
    if not d: return "—"
    try: return datetime.fromisoformat(d).strftime("%d.%m.%Y")
    except: return d

templates.env.globals["fmt"] = fmt
templates.env.globals["ORGANIZATIONS"] = ORGANIZATIONS
templates.env.globals["ORG_MAP"] = ORG_MAP

# ── Auth helpers ───────────────────────────────────────────
def get_session(request: Request):
    token = request.cookies.get("session")
    if not token: return None
    try: return serializer.loads(token, max_age=SESSION_MAX_AGE)
    except BadSignature: return None

def can_edit(session: dict, org_id: int) -> bool:
    """User can edit if admin or curator of this org."""
    if not session: return False
    if session.get("role") == "admin": return True
    allowed = ORG_BY_USERNAME.get(session.get("username"), [])
    return org_id in allowed

def can_edit_kpd(session: dict, employee: str) -> bool:
    """Admin can edit all KPD. Users can edit only their own and ДКВ."""
    if not session: return False
    if session.get("role") == "admin": return True
    if employee == "ДКВ": return True
    username = session.get("username", "")
    _kpd_name = {"akbota": "Акбота", "dinara": "Динара", "ilmira": "Ильмира",
                 "zhanna": "Жанна", "ardak": "Ардак", "chingiz": "Чингиз"}
    my_kpd_name = _kpd_name.get(username.lower(), username)
    return employee == my_kpd_name

templates.env.globals["can_edit"] = can_edit
templates.env.globals["can_edit_kpd"] = can_edit_kpd

def require_session(request: Request):
    s = get_session(request)
    if not s: raise HTTPException(302, headers={"Location": "/login"})
    return s

def redir(url): return RedirectResponse(url, status_code=302)

# ── Login / Logout ─────────────────────────────────────────
@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request, error: str = ""):
    return templates.TemplateResponse("login.html", {"request": request, "error": error})

@app.post("/login")
async def login(request: Request, username: str = Form(...), password: str = Form(...)):
    db = get_db()
    row = db.execute("SELECT * FROM users WHERE username=? AND is_active=1", (username,)).fetchone()
    db.close()
    if not row or not bcrypt.checkpw(password.encode(), row["password_hash"].encode()):
        return redir("/login?error=Неверный+логин+или+пароль")
    token = serializer.dumps({"uid": row["id"], "username": row["username"],
                               "role": row["role"], "full_name": row["full_name"]})
    resp = redir("/")
    resp.set_cookie("session", token, httponly=True, max_age=SESSION_MAX_AGE)
    return resp

@app.get("/logout")
async def logout():
    resp = redir("/login")
    resp.delete_cookie("session")
    return resp

# ── Смена пароля пользователем ─────────────────────────────
@app.get("/change-password", response_class=HTMLResponse)
async def change_password_page(request: Request, error: str = "", success: str = ""):
    s = get_session(request)
    if not s: return redir("/login")
    return templates.TemplateResponse("change_password.html", {
        "request": request, "session": s, "error": error, "success": success
    })

@app.post("/change-password")
async def change_password(
    request: Request,
    current_password: str = Form(...),
    new_password: str = Form(...),
    confirm_password: str = Form(...),
):
    s = get_session(request)
    if not s: return redir("/login")
    uid = s.get("uid")
    db = get_db()
    row = db.execute("SELECT * FROM users WHERE id=?", (uid,)).fetchone()
    if not row:
        return redir("/change-password?error=Пользователь+не+найден")
    if not bcrypt.checkpw(current_password.encode(), row["password_hash"].encode()):
        return redir("/change-password?error=Неверный+текущий+пароль")
    if len(new_password) < 6:
        return redir("/change-password?error=Новый+пароль+должен+быть+не+менее+6+символов")
    if new_password != confirm_password:
        return redir("/change-password?error=Пароли+не+совпадают")
    pw_hash = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode()
    db.execute("UPDATE users SET password_hash=? WHERE id=?", (pw_hash, uid))
    db.commit()
    return redir("/change-password?success=Пароль+успешно+изменён")

# ── Dashboard ──────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def dashboard(request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    stats = []
    for org in ORGANIZATIONS:
        oid = org["id"]
        sessions_count = db.execute("SELECT COUNT(*) FROM sd_sessions WHERE org_id=?", (oid,)).fetchone()[0]
        agenda_count = db.execute("SELECT COUNT(*) FROM sd_agenda_items WHERE session_id IN (SELECT id FROM sd_sessions WHERE org_id=?)", (oid,)).fetchone()[0]
        stats.append({
            "org": org,
            "members":  db.execute("SELECT COUNT(*) FROM sd_members WHERE org_id=?", (oid,)).fetchone()[0],
            "independent": db.execute("SELECT COUNT(*) FROM sd_members WHERE org_id=? AND is_independent=1", (oid,)).fetchone()[0],
            "sessions": sessions_count,
            "agenda":   agenda_count,
            "cmts":     db.execute("SELECT COUNT(*) FROM committees WHERE org_id=?", (oid,)).fetchone()[0],
            "docs":     db.execute("SELECT COUNT(*) FROM documents WHERE org_id=?", (oid,)).fetchone()[0],
        })
    # Ближайшие заседания СД — следующие 30 дней
    from datetime import timedelta
    today_str = datetime.utcnow().strftime("%Y-%m-%d")
    future_str = (datetime.utcnow() + timedelta(days=30)).strftime("%Y-%m-%d")
    upcoming_rows = db.execute(
        "SELECT * FROM sd_sessions WHERE session_date >= ? AND session_date <= ? ORDER BY session_date ASC LIMIT 20",
        (today_str, future_str)
    ).fetchall()
    upcoming = []
    for r in upcoming_rows:
        org = ORG_MAP.get(r["org_id"])
        if org:
            try:
                d = datetime.fromisoformat(r["session_date"])
                delta = (d - datetime.utcnow()).days
                tag = "today" if delta == 0 else ("soon" if 1 <= delta <= 7 else "")
            except Exception:
                delta, tag = 99, ""
            upcoming.append({"org": org, "date": r["session_date"], "format": r["format"], "order_type": r["order_type"], "tag": tag})
    db.close()
    logo_map = {"АКК":"akk","БРК":"brk","ФРП":"frp","ЭКА":"eka","QIC":"qic",
                "Даму":"damu","Отбасы":"otbasy","КЖК":"kjk","КАФ":"kaf","KTD":"ktd"}
    db2 = get_db()
    try:
        agent_unread = db2.execute("SELECT COUNT(*) FROM agent_alerts WHERE is_read=0").fetchone()[0]
    except Exception:
        agent_unread = 0
    db2.close()
    return templates.TemplateResponse("dashboard.html", {"request": request, "session": s, "stats": stats, "logo_map": logo_map, "upcoming": upcoming, "today_str": today_str, "agent_unread": agent_unread})

# ── Org detail page (hub) ──────────────────────────────────
@app.get("/org/{org_id}", response_class=HTMLResponse)
async def org_page(org_id: int, request: Request, tab: str = "sessions", open_sess: int = 0):
    s = get_session(request)
    if not s: return redir("/login")
    if org_id not in ORG_MAP: raise HTTPException(404)
    org = ORG_MAP[org_id]
    editable = can_edit(s, org_id)
    db = get_db()

    # Sessions
    sessions = db.execute("SELECT * FROM sd_sessions WHERE org_id=? ORDER BY session_date DESC", (org_id,)).fetchall()
    agenda_map = {}
    for a in db.execute("SELECT * FROM sd_agenda_items ORDER BY session_id, item_order").fetchall():
        agenda_map.setdefault(a["session_id"], []).append(a)

    # Members
    members = db.execute("SELECT * FROM sd_members WHERE org_id=? ORDER BY id", (org_id,)).fetchall()

    # Committees
    cmts = db.execute("SELECT * FROM committees WHERE org_id=? ORDER BY id", (org_id,)).fetchall()
    cmt_members_map = {}
    for m in db.execute("SELECT * FROM committee_members ORDER BY committee_id, id").fetchall():
        cmt_members_map.setdefault(m["committee_id"], []).append(m)

    cmt_sessions = db.execute("SELECT * FROM committee_sessions WHERE org_id=? ORDER BY session_date ASC", (org_id,)).fetchall()
    participants = db.execute("SELECT * FROM committee_session_participants").fetchall()
    sess_cmt_map = {}
    for p in participants:
        sess_cmt_map.setdefault(p["committee_session_id"], []).append(p["committee_id"])

    sess_members_map = {}
    for m in db.execute("SELECT * FROM committee_session_members").fetchall():
        sess_members_map.setdefault((m["committee_session_id"], m["committee_id"]), []).append(m)

    sess_agenda_map = {}
    for a in db.execute("SELECT * FROM committee_agenda_items ORDER BY committee_session_id, committee_id, item_order").fetchall():
        sess_agenda_map.setdefault((a["committee_session_id"], a["committee_id"]), []).append(a)

    cmt_dict = {c["id"]: dict(c) for c in cmts}

    # Accountable
    accountable = db.execute("SELECT * FROM accountable WHERE org_id=? ORDER BY org_name, position DESC, id", (org_id,)).fetchall()

    # Board
    board = db.execute("SELECT * FROM board_members WHERE org_id=? ORDER BY id", (org_id,)).fetchall()

    # Documents
    docs = db.execute("SELECT id,name,doc_type,date_approved,decision,file_name,file_mime FROM documents WHERE org_id=? ORDER BY created_at DESC", (org_id,)).fetchall()

    # Session files
    sess_files_map = {}
    for f in db.execute("SELECT * FROM sd_session_files").fetchall():
        sess_files_map.setdefault(f["session_id"], []).append(f)

    # Cmt session files
    cmt_files_map = {}
    for f in db.execute("SELECT * FROM cmt_session_files").fetchall():
        cmt_files_map.setdefault(f["committee_session_id"], []).append(f)

    # Work plan
    workplan = db.execute("SELECT * FROM sd_work_plan WHERE org_id=?", (org_id,)).fetchone()

    # EA decisions
    ea_decisions = db.execute("SELECT id,title,decision_date,question_text,requisites,file_name,file_mime,sent_letter,created_at FROM ea_decisions WHERE org_id=? ORDER BY created_at DESC", (org_id,)).fetchall()

    db.close()

    logo_map = {"АКК":"akk","БРК":"brk","ФРП":"frp","ЭКА":"eka","QIC":"qic",
                "Даму":"damu","Отбасы":"otbasy","КЖК":"kjk","КАФ":"kaf","KTD":"ktd"}

    return templates.TemplateResponse("org.html", {
        "request": request, "session": s, "org": org, "tab": tab,
        "open_sess": open_sess,
        "editable": editable,
        "sessions": sessions, "agenda_map": agenda_map,
        "sess_files_map": sess_files_map,
        "members": members,
        "committees": cmts, "cmt_members_map": cmt_members_map,
        "cmt_sessions": cmt_sessions, "sess_cmt_map": sess_cmt_map,
        "sess_members_map": sess_members_map, "sess_agenda_map": sess_agenda_map,
        "cmt_dict": cmt_dict, "cmt_files_map": cmt_files_map,
        "accountable": accountable,
        "board": board,
        "docs": docs,
        "workplan": workplan,
        "ea_decisions": ea_decisions,
        "org_logo": logo_map.get(org["short"], ""),
    })

# ── Guard helper ───────────────────────────────────────────
def guard(request, org_id):
    s = get_session(request)
    if not s: return redir("/login"), None
    if not can_edit(s, org_id):
        return redir(f"/org/{org_id}?error=noaccess"), None
    return None, s

# ── SD Sessions ────────────────────────────────────────────
@app.post("/org/{org_id}/sessions/add")
async def add_session(org_id: int, request: Request,
    session_date: str = Form(...), format: str = Form("Очное"), order_type: str = Form("Очередное")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("INSERT INTO sd_sessions (org_id,session_date,format,order_type,created_at,updated_at) VALUES (?,?,?,?,?,?)",
        (org_id, session_date, format, order_type, now(), now()))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=sessions")

@app.post("/org/{org_id}/sessions/delete/{sid}")
async def del_session(org_id: int, sid: int, request: Request):
    err, s = guard(request, org_id)
    if err: return err
    if s["role"] != "admin": return redir(f"/org/{org_id}?tab=sessions")
    db = get_db()
    db.execute("DELETE FROM sd_sessions WHERE id=? AND org_id=?", (sid, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=sessions")

@app.post("/org/{org_id}/sessions/edit/{sid}")
async def edit_session(org_id: int, sid: int, request: Request,
    open_sess: int = Query(0),
    session_date: str = Form(""), format: str = Form("Очное"),
    order_type: str = Form("Очередное")):
    err, s = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("UPDATE sd_sessions SET session_date=?, format=?, order_type=? WHERE id=? AND org_id=?",
        (session_date, format, order_type, sid, org_id))
    db.commit(); db.close()
    log_activity(s, "edit_session", org_id, f"{session_date} {format} {order_type}")
    return redir(f"/org/{org_id}?tab=sessions&open_sess={sid}")

@app.post("/org/{org_id}/sessions/{sid}/agenda/add")
async def add_agenda(org_id: int, sid: int, request: Request,
    text: str = Form(""), status: str = Form("done"), note: str = Form("")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    cnt = db.execute("SELECT COUNT(*) FROM sd_agenda_items WHERE session_id=?", (sid,)).fetchone()[0]
    db.execute("INSERT INTO sd_agenda_items (session_id,item_order,text,status,note) VALUES (?,?,?,?,?)",
        (sid, cnt+1, text, status, note))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=sessions&open_sess={sid}")

@app.post("/org/{org_id}/sessions/{sid}/agenda/edit/{aid}")
async def edit_agenda(org_id: int, sid: int, aid: int, request: Request,
    text: str = Form(""), status: str = Form("done"), note: str = Form("")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("UPDATE sd_agenda_items SET text=?,status=?,note=? WHERE id=? AND session_id=?",
        (text, status, note, aid, sid))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=sessions&open_sess={sid}")

@app.post("/org/{org_id}/sessions/{sid}/agenda/delete/{aid}")
async def del_agenda(org_id: int, sid: int, aid: int, request: Request):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("DELETE FROM sd_agenda_items WHERE id=? AND session_id=?", (aid, sid))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=sessions&open_sess={sid}")

# ── SD Members ─────────────────────────────────────────────
@app.post("/org/{org_id}/members/add")
async def add_member(org_id: int, request: Request, full_name: str = Form(...),
    role: str = Form("mem"), is_independent: int = Form(0), position: str = Form(""),
    date_from: str = Form(""), date_to: str = Form(""), decision: str = Form("")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("INSERT INTO sd_members (org_id,full_name,role,is_independent,position,date_from,date_to,decision,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
        (org_id, full_name, role, is_independent, position, date_from, date_to, decision, now(), now()))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=members")

@app.post("/org/{org_id}/members/edit/{mid}")
async def edit_member(org_id: int, mid: int, request: Request, full_name: str = Form(...),
    role: str = Form("mem"), is_independent: int = Form(0), position: str = Form(""),
    date_from: str = Form(""), date_to: str = Form(""), decision: str = Form("")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("UPDATE sd_members SET full_name=?,role=?,is_independent=?,position=?,date_from=?,date_to=?,decision=?,updated_at=? WHERE id=? AND org_id=?",
        (full_name, role, is_independent, position, date_from, date_to, decision, now(), mid, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=members")

@app.post("/org/{org_id}/members/delete/{mid}")
async def del_member(org_id: int, mid: int, request: Request):
    err, s = guard(request, org_id)
    if err: return err
    if s["role"] != "admin": return redir(f"/org/{org_id}?tab=members")
    db = get_db()
    db.execute("DELETE FROM sd_members WHERE id=? AND org_id=?", (mid, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=members")

# ── Committees ─────────────────────────────────────────────
@app.post("/org/{org_id}/committees/add")
async def add_cmt(org_id: int, request: Request, name: str = Form(...), color: str = Form("#1a6b3c")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("INSERT INTO committees (org_id,name,color,created_at) VALUES (?,?,?,?)", (org_id, name, color, now()))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees")

@app.post("/org/{org_id}/committees/edit/{cid}")
async def edit_cmt(org_id: int, cid: int, request: Request, name: str = Form(...), color: str = Form("#1a6b3c")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("UPDATE committees SET name=?, color=? WHERE id=? AND org_id=?", (name, color, cid, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees")

@app.post("/org/{org_id}/committees/{cid}/members/edit/{mid}")
async def edit_cmt_member(org_id: int, cid: int, mid: int, request: Request,
    full_name: str = Form(...), role: str = Form("Член комитета")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("UPDATE committee_members SET full_name=?, role=? WHERE id=? AND committee_id=?", (full_name, role, mid, cid))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees")


@app.post("/org/{org_id}/committees/delete/{cid}")
async def del_cmt(org_id: int, cid: int, request: Request):
    err, s = guard(request, org_id)
    if err: return err
    if s["role"] != "admin": return redir(f"/org/{org_id}?tab=committees")
    db = get_db()
    db.execute("DELETE FROM committees WHERE id=? AND org_id=?", (cid, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees")

@app.post("/org/{org_id}/committees/{cid}/members/add")
async def add_cmt_member(org_id: int, cid: int, request: Request,
    full_name: str = Form(...), role: str = Form("Член комитета")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("INSERT INTO committee_members (committee_id,full_name,role) VALUES (?,?,?)", (cid, full_name, role))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees")

@app.post("/org/{org_id}/committees/{cid}/members/delete/{mid}")
async def del_cmt_member(org_id: int, cid: int, mid: int, request: Request):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("DELETE FROM committee_members WHERE id=? AND committee_id=?", (mid, cid))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees")

@app.post("/org/{org_id}/cmt-sessions/add")
async def add_cmt_sess(org_id: int, request: Request,
    session_date: str = Form(...), protocol_num: str = Form(""), committee_ids: list = Form(...)):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    cur = db.execute("INSERT INTO committee_sessions (org_id,session_date,protocol_num,created_at) VALUES (?,?,?,?)",
        (org_id, session_date, protocol_num, now()))
    sid = cur.lastrowid
    for cid in committee_ids:
        cid = int(cid)
        db.execute("INSERT INTO committee_session_participants (committee_session_id,committee_id) VALUES (?,?)", (sid, cid))
        for m in db.execute("SELECT * FROM committee_members WHERE committee_id=?", (cid,)).fetchall():
            db.execute("INSERT INTO committee_session_members (committee_session_id,committee_id,full_name,role) VALUES (?,?,?,?)",
                (sid, cid, m["full_name"], m["role"]))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees&ctab=sessions")

@app.post("/org/{org_id}/cmt-sessions/delete/{sid}")
async def del_cmt_sess(org_id: int, sid: int, request: Request):
    err, s = guard(request, org_id)
    if err: return err
    if s["role"] != "admin": return redir(f"/org/{org_id}?tab=committees&ctab=sessions")
    db = get_db()
    db.execute("DELETE FROM committee_sessions WHERE id=? AND org_id=?", (sid, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees&ctab=sessions")

@app.post("/org/{org_id}/cmt-sessions/{sid}/agenda/add")
async def add_cmt_agenda(org_id: int, sid: int, request: Request,
    committee_id: int = Form(...), text: str = Form(""), status: str = Form("done"), note: str = Form("")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    cnt = db.execute("SELECT COUNT(*) FROM committee_agenda_items WHERE committee_session_id=? AND committee_id=?",
        (sid, committee_id)).fetchone()[0]
    db.execute("INSERT INTO committee_agenda_items (committee_session_id,committee_id,item_order,text,status,note) VALUES (?,?,?,?,?,?)",
        (sid, committee_id, cnt+1, text, status, note))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees&ctab=sessions")

@app.post("/org/{org_id}/cmt-sessions/{sid}/agenda/delete/{aid}")
async def del_cmt_agenda(org_id: int, sid: int, aid: int, request: Request):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("DELETE FROM committee_agenda_items WHERE id=? AND committee_session_id=?", (aid, sid))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees&ctab=sessions")

# ── Accountable ────────────────────────────────────────────
@app.post("/org/{org_id}/accountable/add")
async def add_acc(org_id: int, request: Request, full_name: str = Form(...),
    position: str = Form("emp"), org_name: str = Form(""), phone: str = Form(""),
    email: str = Form(""), date_from: str = Form(""), date_to: str = Form(""), decision: str = Form("")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("INSERT INTO accountable (org_id,full_name,position,org_name,phone,email,date_from,date_to,decision,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
        (org_id, full_name, position, org_name, phone, email, date_from, date_to, decision, now(), now()))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=accountable")

@app.post("/org/{org_id}/accountable/edit/{aid}")
async def edit_acc(org_id: int, aid: int, request: Request, full_name: str = Form(...),
    position: str = Form("emp"), org_name: str = Form(""), phone: str = Form(""),
    email: str = Form(""), date_from: str = Form(""), date_to: str = Form(""), decision: str = Form("")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("UPDATE accountable SET full_name=?,position=?,org_name=?,phone=?,email=?,date_from=?,date_to=?,decision=?,updated_at=? WHERE id=? AND org_id=?",
        (full_name, position, org_name, phone, email, date_from, date_to, decision, now(), aid, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=accountable")

@app.post("/org/{org_id}/accountable/delete/{aid}")
async def del_acc(org_id: int, aid: int, request: Request):
    err, s = guard(request, org_id)
    if err: return err
    if s["role"] != "admin": return redir(f"/org/{org_id}?tab=accountable")
    db = get_db()
    db.execute("DELETE FROM accountable WHERE id=? AND org_id=?", (aid, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=accountable")

# ── Board ──────────────────────────────────────────────────
@app.post("/org/{org_id}/board/add")
async def add_board(org_id: int, request: Request, full_name: str = Form(...),
    position: str = Form(""), date_from: str = Form(""), date_to: str = Form("")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("INSERT INTO board_members (org_id,full_name,position,date_from,date_to,created_at,updated_at) VALUES (?,?,?,?,?,?,?)",
        (org_id, full_name, position, date_from, date_to, now(), now()))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=board")

@app.post("/org/{org_id}/board/edit/{bid}")
async def edit_board(org_id: int, bid: int, request: Request, full_name: str = Form(...),
    position: str = Form(""), date_from: str = Form(""), date_to: str = Form("")):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("UPDATE board_members SET full_name=?,position=?,date_from=?,date_to=?,updated_at=? WHERE id=? AND org_id=?",
        (full_name, position, date_from, date_to, now(), bid, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=board")

@app.post("/org/{org_id}/board/delete/{bid}")
async def del_board(org_id: int, bid: int, request: Request):
    err, s = guard(request, org_id)
    if err: return err
    if s["role"] != "admin": return redir(f"/org/{org_id}?tab=board")
    db = get_db()
    db.execute("DELETE FROM board_members WHERE id=? AND org_id=?", (bid, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=board")

# ── Documents ──────────────────────────────────────────────
@app.post("/org/{org_id}/documents/add")
async def add_doc(org_id: int, request: Request, name: str = Form(...),
    doc_type: str = Form("Положение"), date_approved: str = Form(""), decision: str = Form(""),
    file: UploadFile = File(None)):
    err, _ = guard(request, org_id)
    if err: return err
    fname = fdata = fmime = None
    if file and file.filename:
        fname = file.filename; fmime = file.content_type; fdata = await file.read()
    db = get_db()
    db.execute("INSERT INTO documents (org_id,name,doc_type,date_approved,decision,file_name,file_data,file_mime,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
        (org_id, name, doc_type, date_approved, decision, fname, fdata, fmime, now(), now()))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=documents")

@app.post("/org/{org_id}/documents/delete/{did}")
async def del_doc(org_id: int, did: int, request: Request):
    err, s = guard(request, org_id)
    if err: return err
    if s["role"] != "admin": return redir(f"/org/{org_id}?tab=documents")
    db = get_db()
    db.execute("DELETE FROM documents WHERE id=? AND org_id=?", (did, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=documents")

@app.get("/org/{org_id}/documents/view/{did}")
async def view_doc(org_id: int, did: int, request: Request, download: str = "0"):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    row = db.execute("SELECT file_data,file_mime,file_name FROM documents WHERE id=? AND org_id=?", (did, org_id)).fetchone()
    db.close()
    if not row or not row["file_data"]: raise HTTPException(404)
    from urllib.parse import quote
    disposition = "attachment" if download=="1" else "inline"
    fname = row["file_name"] or "document"
    fname_encoded = quote(fname)
    return StreamingResponse(io.BytesIO(row["file_data"]),
        media_type=row["file_mime"] or "application/octet-stream",
        headers={"Content-Disposition": f"{disposition}; filename*=UTF-8''{fname_encoded}"})

@app.post("/org/{org_id}/documents/edit/{did}")
async def edit_doc(org_id: int, did: int, request: Request,
    name: str = Form(...), doc_type: str = Form("Положение"),
    date_approved: str = Form(""), decision: str = Form(""),
    file: UploadFile = File(None)):
    err, s = guard(request, org_id)
    if err: return err
    db = get_db()
    if file and file.filename:
        fdata = await file.read()
        db.execute("UPDATE documents SET name=?,doc_type=?,date_approved=?,decision=?,file_name=?,file_data=?,file_mime=?,updated_at=? WHERE id=? AND org_id=?",
            (name, doc_type, date_approved, decision, file.filename, fdata, file.content_type, now(), did, org_id))
    else:
        db.execute("UPDATE documents SET name=?,doc_type=?,date_approved=?,decision=?,updated_at=? WHERE id=? AND org_id=?",
            (name, doc_type, date_approved, decision, now(), did, org_id))
    db.commit(); db.close()
    log_activity(s, "edit_document", org_id, name)
    return redir(f"/org/{org_id}?tab=documents")

# ── Users ──────────────────────────────────────────────────
@app.get("/users", response_class=HTMLResponse)
async def users_page(request: Request):
    s = get_session(request)
    if not s or s["role"] != "admin": return redir("/login")
    db = get_db()
    users = db.execute("SELECT id,username,full_name,role,is_active,created_at FROM users ORDER BY id").fetchall()
    db.close()
    db2 = get_db()
    pwd_requests = db2.execute("SELECT * FROM password_requests ORDER BY created_at DESC").fetchall()
    db2.close()
    return templates.TemplateResponse("users.html", {"request": request, "session": s, "users": users, "pwd_requests": pwd_requests})

@app.post("/users/add")
async def add_user(request: Request, username: str = Form(...), full_name: str = Form(...),
    role: str = Form("staff"), password: str = Form(...)):
    s = get_session(request)
    if not s or s["role"] != "admin": return redir("/login")
    pw = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()
    db = get_db()
    try:
        db.execute("INSERT INTO users (username,full_name,password_hash,role,is_active,created_at) VALUES (?,?,?,?,1,?)",
            (username, full_name, pw, role, now()))
        db.commit()
    except: pass
    db.close()
    return redir("/users")

@app.post("/users/toggle/{uid}")
async def toggle_user(uid: int, request: Request):
    s = get_session(request)
    if not s or s["role"] != "admin": return redir("/login")
    db = get_db()
    row = db.execute("SELECT is_active FROM users WHERE id=?", (uid,)).fetchone()
    if row:
        db.execute("UPDATE users SET is_active=? WHERE id=?", (0 if row["is_active"] else 1, uid))
        db.commit()
    db.close()
    return redir("/users")

@app.post("/users/delete/{uid}")
async def delete_user(uid: int, request: Request):
    s = get_session(request)
    if not s or s["role"] != "admin": return redir("/login")
    db = get_db()
    db.execute("DELETE FROM users WHERE id=? AND username != 'admin'", (uid,))
    db.commit(); db.close()
    return redir("/users")

@app.post("/users/reset-password/{uid}")
async def reset_pwd(uid: int, request: Request, new_password: str = Form(...)):
    s = get_session(request)
    if not s or s["role"] != "admin": return redir("/login")
    pw = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode()
    db = get_db()
    db.execute("UPDATE users SET password_hash=? WHERE id=?", (pw, uid))
    db.commit(); db.close()
    return redir("/users")

if __name__ == "__main__":
    init_db()
    from database import migrate_db
    migrate_db()
    print("\n" + "="*60)
    print("  АО «НИХ» Байтерек — Корпоративная карточка ДО/ДЗО")
    print("  Браузер: http://localhost:8000")
    print("  Логины: admin/admin123, chingiz/pass123,")
    print("          akbota/pass123, dinara/pass123,")
    print("          ilmira/pass123, zhanna/pass123")
    print("="*60 + "\n")
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=False)

# ── Search API ─────────────────────────────────────────────
@app.get("/org/{org_id}/search")
async def search_agenda(org_id: int, request: Request, q: str = "", tab: str = "sessions"):
    s = get_session(request)
    if not s: return JSONResponse({"results": [], "q": q, "count": 0})
    if not q or len(q.strip()) < 1: return JSONResponse({"results": [], "q": q, "count": 0})
    
    db = get_db()
    query = f"%{q.strip()}%"
    results = []
    
    if tab == "sessions":
        rows = db.execute("""
            SELECT a.id, a.text, a.status, a.note, a.item_order,
                   s.id as sess_id, s.session_date, s.format, s.order_type
            FROM sd_agenda_items a
            JOIN sd_sessions s ON s.id = a.session_id
            WHERE s.org_id = ? AND LOWER(a.text) LIKE LOWER(?)
            ORDER BY s.session_date ASC, a.item_order ASC
        """, (org_id, query)).fetchall()
        
        for i, r in enumerate(rows):
            results.append({
                "type": "sess",
                "sess_id": r["sess_id"],
                "date": fmt(r["session_date"]),
                "format": r["format"],
                "order": r["order_type"],
                "q_num": r["item_order"] or (i+1),
                "text": r["text"] or "",
                "status": r["status"] or "notdone",
                "note": r["note"] or ""
            })
    
    elif tab == "committees":
        rows = db.execute("""
            SELECT a.text, a.status, a.item_order, a.committee_id,
                   cs.id as sess_id, cs.session_date, cs.protocol_num,
                   c.name as cmt_name
            FROM committee_agenda_items a
            JOIN committee_sessions cs ON cs.id = a.committee_session_id
            JOIN committees c ON c.id = a.committee_id
            WHERE cs.org_id = ? AND LOWER(a.text) LIKE LOWER(?)
            ORDER BY cs.session_date ASC, a.item_order ASC
        """, (org_id, query)).fetchall()
        
        for i, r in enumerate(rows):
            results.append({
                "type": "cmt",
                "sess_id": r["sess_id"],
                "date": fmt(r["session_date"]),
                "cmt": r["cmt_name"] or "",
                "proto": r["protocol_num"] or "",
                "q_num": r["item_order"] or (i+1),
                "text": r["text"] or "",
                "status": r["status"] or "notdone"
            })
    
    db.close()
    return JSONResponse({"results": results, "q": q, "count": len(results)})

# ── Auto-Update Panel ─────────────────────────────────────
import urllib.request, json as _json

@app.get("/admin/update", response_class=HTMLResponse)
async def update_panel(request: Request):
    s = get_session(request)
    if not s or s["role"] != "admin": return redir("/login")
    html = """<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="UTF-8">
<title>Обновление системы</title>
<link href="https://fonts.googleapis.com/css2?family=Montserrat:wght@400;600;700;800&display=swap" rel="stylesheet">
<style>
*{box-sizing:border-box;margin:0;padding:0;}
body{font-family:'Montserrat',sans-serif;background:#f5f0e8;padding:30px;color:#2d3a1f;}
h1{font-size:18px;font-weight:900;margin-bottom:6px;}
.sub{font-size:11px;color:#8a9a70;margin-bottom:24px;}
.card{background:#fff;border:1px solid #d4cdb8;border-radius:8px;padding:20px;margin-bottom:14px;}
.card h2{font-size:13px;font-weight:700;margin-bottom:6px;}
.card p{font-size:11px;color:#8a9a70;margin-bottom:12px;}
.btn{display:inline-block;padding:9px 20px;background:#2d3a1f;color:#f5f0e8;border:none;border-radius:4px;font-family:inherit;font-size:11px;font-weight:700;cursor:pointer;text-decoration:none;letter-spacing:.5px;}
.btn:hover{background:#7ab317;color:#2d3a1f;}
.btn-red{background:#c62828;}
.result{margin-top:12px;padding:10px 14px;border-radius:4px;font-size:11px;display:none;}
.result.ok{background:#e8f5e9;color:#2e7d32;border:1px solid #a5d6a7;}
.result.err{background:#ffebee;color:#c62828;border:1px solid #ef9a9a;}
a.back{display:inline-block;margin-bottom:16px;font-size:11px;color:#8a9a70;text-decoration:none;}
a.back:hover{color:#2d3a1f;}
</style>
</head>
<body>
<a class="back" href="/">← На главную</a>
<h1>⚙️ Панель обновлений</h1>
<p class="sub">Применяй обновления прямо в браузере — без замены файлов</p>

<div id="updates"></div>

<script>
const UPDATES = [
  {id:'fix2', name:'Заседание остаётся открытым', desc:'После добавления вопроса повестки заседание не сворачивается', endpoint:'/admin/apply/fix2'},
  {id:'fix3', name:'Документы: просмотр и скачивание', desc:'Кнопки «Просмотр» и «Скачать» + редактирование документов', endpoint:'/admin/apply/fix3'},
  {id:'fix4', name:'ЕА: загрузка файла', desc:'Файл прикрепляется кликом и перетаскиванием', endpoint:'/admin/apply/fix4'},
  {id:'fix5', name:'Календарь: добавление событий', desc:'Кнопка + Добавить событие в календаре', endpoint:'/admin/apply/fix5'},
  {id:'fix6', name:'Комитеты: редактирование', desc:'Кнопки редактирования заседаний и состава комитетов', endpoint:'/admin/apply/fix6'},
  {id:'fix7', name:'Логотип Байтерека', desc:'Логотип в навигации вместо текста НИХ', endpoint:'/admin/apply/fix7'},
  {id:'fix8', name:'Тёмная тема: цвета текста', desc:'Исправлены цвета в тёмной теме', endpoint:'/admin/apply/fix8'},
  {id:'fix9', name:'Активность: компактный онлайн', desc:'Онлайн-пользователи отображаются в одну строку', endpoint:'/admin/apply/fix9'},
];

const cont = document.getElementById('updates');
UPDATES.forEach(u => {
  cont.innerHTML += `<div class="card">
    <h2>${u.name}</h2>
    <p>${u.desc}</p>
    <button class="btn" onclick="apply('${u.id}','${u.endpoint}')">▶ Применить</button>
    <div class="result" id="res_${u.id}"></div>
  </div>`;
});

function apply(id, endpoint){
  const res = document.getElementById('res_' + id);
  res.style.display='block';
  res.className='result';
  res.textContent='⏳ Применяем...';
  fetch(endpoint, {method:'POST'})
    .then(r => r.json())
    .then(d => {
      res.className = 'result ' + (d.ok ? 'ok' : 'err');
      res.textContent = d.ok ? '✅ ' + d.msg : '❌ ' + d.msg;
    })
    .catch(e => { res.className='result err'; res.textContent='❌ Ошибка: ' + e; });
}
</script>
</body>
</html>"""
    return HTMLResponse(html)

@app.post("/admin/apply/{fix_id}")
async def apply_fix(fix_id: str, request: Request):
    s = get_session(request)
    if not s or s["role"] != "admin":
        return JSONResponse({"ok": False, "msg": "Нет доступа"})
    
    import os
    base = os.path.dirname(os.path.abspath(__file__))
    tmpl = os.path.join(base, "templates")
    
    try:
        if fix_id == "fix2":
            with open(os.path.join(base, "main.py"), "r", encoding="utf-8") as f:
                c = f.read()
            old = '    return redir(f"/org/{org_id}?tab=sessions")\n\n@app.post("/org/{org_id}/sessions/{sid}/agenda/edit'
            new = '    return redir(f"/org/{org_id}?tab=sessions&open_sess={sid}")\n\n@app.post("/org/{org_id}/sessions/{sid}/agenda/edit'
            if old not in c:
                return JSONResponse({"ok": True, "msg": "Уже применено"})
            c = c.replace(old, new)
            with open(os.path.join(base, "main.py"), "w", encoding="utf-8") as f:
                f.write(c)
            return JSONResponse({"ok": True, "msg": "Готово! Перезапустите сервер."})
        
        elif fix_id == "fix9":
            activity = """{% extends "base.html" %}
{% set active = "activity" %}
{% block title %}Активность — НИХ Байтерек{% endblock %}
{% block content %}
<div class="page-hdr">
  <div class="page-title">📊 Активность</div>
</div>
<div style="font-size:9px;font-weight:700;color:var(--mid);text-transform:uppercase;letter-spacing:2px;margin-bottom:10px;">🟢 Сейчас онлайн</div>
<div class="card" style="margin-bottom:20px;">
  <div style="display:flex;flex-wrap:wrap;gap:8px;padding:12px 16px;">
    {% for u in online %}
    <div style="display:inline-flex;align-items:center;gap:6px;background:var(--cream2);border:1px solid var(--light);border-radius:20px;padding:5px 14px;">
      <span style="color:#2e7d32;font-size:11px;">●</span>
      <span style="font-size:11px;font-weight:700;">{{ u.full_name }}</span>
      <span style="font-size:10px;color:var(--mid);">{{ u.username }}</span>
    </div>
    {% else %}
    <span style="color:var(--mid);font-size:12px;">Нет активных пользователей</span>
    {% endfor %}
  </div>
</div>
<div style="font-size:9px;font-weight:700;color:var(--mid);text-transform:uppercase;letter-spacing:2px;margin-bottom:10px;">📋 История изменений</div>
<div class="card">
  <table class="tbl">
    <thead><tr><th>Дата / Время</th><th>Пользователь</th><th>Действие</th><th>Организация</th><th>Детали</th></tr></thead>
    <tbody>
    {% for log in logs %}
    <tr>
      <td style="font-size:10px;color:var(--mid);white-space:nowrap;">{{ log.created_at[:16].replace("T"," ") }}</td>
      <td style="font-weight:600;font-size:12px;">{{ log.full_name }}<br><span style="font-size:9px;color:var(--mid);">{{ log.username }}</span></td>
      <td>
        {% if "upload" in log.action %}<span class="badge" style="background:#f3e5f5;color:#7b1fa2;border:1px solid #ce93d8;">⬆ Загрузка</span>
        {% elif "delete" in log.action %}<span class="badge b-notdone">✕ Удаление</span>
        {% elif "edit" in log.action %}<span class="badge" style="background:#e3f2fd;color:#1565c0;border:1px solid #90caf9;">✎ Изменение</span>
        {% elif "add" in log.action %}<span class="badge b-done">+ Добавление</span>
        {% else %}<span class="badge b-mem">{{ log.action }}</span>{% endif %}
      </td>
      <td style="font-size:11px;font-weight:700;">
        {% if log.org_id %}{% for org in ORGANIZATIONS %}{% if org.id==log.org_id %}{{ org.short }}{% endif %}{% endfor %}{% else %}—{% endif %}
      </td>
      <td style="font-size:11px;color:var(--mid);max-width:200px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;">{{ log.details or "—" }}</td>
    </tr>
    {% else %}
    <tr><td colspan="5" style="text-align:center;padding:20px;color:var(--mid);">История пуста</td></tr>
    {% endfor %}
    </tbody>
  </table>
</div>
{% endblock %}"""
            with open(os.path.join(tmpl, "activity.html"), "w", encoding="utf-8") as f:
                f.write(activity)
            return JSONResponse({"ok": True, "msg": "Готово! Обновите страницу активности."})
        
        else:
            return JSONResponse({"ok": False, "msg": f"Обновление {fix_id} пока не реализовано"})
    
    except Exception as e:
        return JSONResponse({"ok": False, "msg": str(e)})


# ── DKV Documents ─────────────────────────────────────────────
@app.get("/dkv/documents", response_class=HTMLResponse)
async def dkv_docs_page(request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    docs = db.execute("SELECT * FROM dkv_documents ORDER BY created_at DESC").fetchall()
    db.close()
    return templates.TemplateResponse("dkv_documents.html", {"request": request, "session": s, "docs": docs})

@app.post("/dkv/documents/add")
async def dkv_doc_add(request: Request, name: str = Form(...), file: UploadFile = File(None)):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    fdata = await file.read() if file and file.filename else None
    db.execute("INSERT INTO dkv_documents (name,file_name,file_data,file_mime,uploaded_by,created_at) VALUES (?,?,?,?,?,?)",
        (name, file.filename if file else None, fdata, file.content_type if file else None, s["full_name"], now()))
    db.commit(); db.close()
    return redir("/dkv/documents")

@app.get("/dkv/documents/view/{did}")
async def dkv_doc_view(did: int, request: Request, download: str = "0"):
    from urllib.parse import quote
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    row = db.execute("SELECT * FROM dkv_documents WHERE id=?", (did,)).fetchone()
    db.close()
    if not row or not row["file_data"]: raise HTTPException(404)
    disposition = "attachment" if download == "1" else "inline"
    fname_encoded = quote(row["file_name"] or "document")
    return StreamingResponse(io.BytesIO(row["file_data"]),
        media_type=row["file_mime"] or "application/octet-stream",
        headers={"Content-Disposition": f"{disposition}; filename*=UTF-8''{fname_encoded}"})

@app.post("/dkv/documents/delete/{did}")
async def dkv_doc_delete(did: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    db.execute("DELETE FROM dkv_documents WHERE id=?", (did,))
    db.commit(); db.close()
    return redir("/dkv/documents")

@app.get("/dkv/documents/edit/{did}", response_class=HTMLResponse)
async def dkv_doc_edit_page(did: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    doc = db.execute("SELECT * FROM dkv_documents WHERE id=?", (did,)).fetchone()
    db.close()
    if not doc: raise HTTPException(404)
    return templates.TemplateResponse("dkv_documents_edit.html", {
        "request": request, "session": s, "doc": doc
    })

@app.post("/dkv/documents/update/{did}")
async def dkv_doc_update(did: int, request: Request, name: str = Form(...), file: UploadFile = File(None)):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    doc = db.execute("SELECT * FROM dkv_documents WHERE id=?", (did,)).fetchone()
    if not doc: raise HTTPException(404)
    
    # Если новый файл загружен, используем его; иначе сохраняем старый
    fdata = await file.read() if file and file.filename else doc["file_data"]
    fname = file.filename if file and file.filename else doc["file_name"]
    fmime = file.content_type if file and file.filename else doc["file_mime"]
    
    db.execute("UPDATE dkv_documents SET name=?, file_name=?, file_data=?, file_mime=? WHERE id=?",
        (name, fname, fdata, fmime, did))
    db.commit(); db.close()
    return redir("/dkv/documents")

# ── DKV KPD ───────────────────────────────────────────────────
@app.get("/dkv/kpd", response_class=HTMLResponse)
async def dkv_kpd_page(request: Request, tab: str = "ДКВ", view: str = "grid", quarter: str = "I"):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    # Фильтруем по employee и quarter
    kpd_rows = db.execute(
        "SELECT * FROM dkv_kpd WHERE employee=? AND quarter=? ORDER BY item_order",
        (tab, quarter)
    ).fetchall()
    db.close()
    employees = ["ДКВ", "Ардак", "Ильмира", "Акбота", "Жанна", "Динара", "Чингиз"]
    quarters = ["I", "II", "III", "IV"]
    
    # Если view == "grid" — показываем красивую визуализацию
    if view == "grid":
        return templates.TemplateResponse("dkv_kpd_grid.html", {
            "request": request, "session": s,
            "kpd_rows": kpd_rows, "tab": tab, "quarter": quarter, 
            "employees": employees, "quarters": quarters
        })
    
    # Иначе — обычную таблицу (старый вид)
    return templates.TemplateResponse("dkv_kpd.html", {
        "request": request, "session": s,
        "kpd_rows": kpd_rows, "tab": tab, "quarter": quarter, "employees": employees, "quarters": quarters
    })

@app.post("/dkv/kpd/add")
async def dkv_kpd_add(request: Request,
    employee: str = Form("map"),
    quarter: str = Form("I"),
    name: str = Form(""),
    unit: str = Form(""),
    formula: str = Form(""),
    formula_desc: str = Form(""),
    data_source: str = Form(""),
    weight: str = Form(""),
    value_type: str = Form(""),
    threshold: str = Form(""),
    goal: str = Form(""),
    challenge: str = Form(""),
    fact: str = Form(""),
    file: UploadFile = File(None)):
    s = get_session(request)
    if not s: return redir("/login")
    
    # Проверка прав: admin — всё, обычный пользователь — только своё и ДКВ
    role = s.get("role")
    username = s.get("username", "")
    # Маппинг username (латиница) → имя в КПД (кириллица)
    _kpd_name = {"akbota": "Акбота", "dinara": "Динара", "ilmira": "Ильмира",
                 "zhanna": "Жанна", "ardak": "Ардак", "chingiz": "Чингиз"}
    my_kpd_name = _kpd_name.get(username.lower(), username)
    if role != "admin" and employee != "ДКВ" and my_kpd_name != employee:
        return redir(f"/dkv/kpd?tab={employee}&quarter={quarter}&error=Нет+прав")
    
    db = get_db()
    cnt = db.execute("SELECT COUNT(*) FROM dkv_kpd WHERE employee=? AND quarter=?", (employee, quarter)).fetchone()[0]
    fdata = await file.read() if file and file.filename else None
    db.execute("""INSERT INTO dkv_kpd 
        (employee,quarter,item_order,name,unit,formula,formula_desc,data_source,weight,value_type,threshold,goal,challenge,fact,file_name,file_data,file_mime,created_at)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (employee, quarter, cnt+1, name, unit, formula, formula_desc, data_source, weight, value_type,
         threshold, goal, challenge, fact,
         file.filename if file and file.filename else None, fdata,
         file.content_type if file and file.filename else None, now()))
    db.commit(); db.close()
    return redir(f"/dkv/kpd?tab={employee}&quarter={quarter}")

@app.post("/dkv/kpd/delete/{kid}")
async def dkv_kpd_delete(kid: int, request: Request, employee: str = Form("ДКВ"), quarter: str = Form("I")):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    row = db.execute("SELECT * FROM dkv_kpd WHERE id=?", (kid,)).fetchone()
    if not row:
        db.close(); raise HTTPException(404)
    emp = row["employee"]
    qtr = row["quarter"]
    role = s.get("role"); username = s.get("username", "")
    _kpd_name = {"akbota": "Акбота", "dinara": "Динара", "ilmira": "Ильмира",
                 "zhanna": "Жанна", "ardak": "Ардак", "chingiz": "Чингиз"}
    my_kpd_name = _kpd_name.get(username.lower(), username)
    if role != "admin" and emp != "ДКВ" and my_kpd_name != emp:
        db.close(); return redir(f"/dkv/kpd?tab={emp}&quarter={qtr}&error=Нет+прав")
    db.execute("DELETE FROM dkv_kpd WHERE id=?", (kid,))
    db.commit(); db.close()
    return redir(f"/dkv/kpd?tab={emp}&quarter={qtr}")

@app.get("/dkv/kpd/edit/{kid}", response_class=HTMLResponse)
async def dkv_kpd_edit_page(kid: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    row = db.execute("SELECT * FROM dkv_kpd WHERE id=?", (kid,)).fetchone()
    db.close()
    if not row: raise HTTPException(404)
    return templates.TemplateResponse("dkv_kpd_edit.html", {
        "request": request, "session": s, "row": row
    })

@app.post("/dkv/kpd/update/{kid}")
async def dkv_kpd_update(kid: int, request: Request,
    name: str = Form(""),
    unit: str = Form(""),
    formula: str = Form(""),
    formula_desc: str = Form(""),
    data_source: str = Form(""),
    weight: str = Form(""),
    value_type: str = Form(""),
    threshold: str = Form(""),
    goal: str = Form(""),
    challenge: str = Form(""),
    fact: str = Form(""),
    file: UploadFile = File(None)):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    row = db.execute("SELECT * FROM dkv_kpd WHERE id=?", (kid,)).fetchone()
    if not row:
        db.close(); raise HTTPException(404)
    emp = row["employee"]
    role = s.get("role"); username = s.get("username", "")
    _kpd_name = {"akbota": "Акбота", "dinara": "Динара", "ilmira": "Ильмира",
                 "zhanna": "Жанна", "ardak": "Ардак", "chingiz": "Чингиз"}
    my_kpd_name = _kpd_name.get(username.lower(), username)
    if role != "admin" and emp != "ДКВ" and my_kpd_name != emp:
        db.close(); return redir(f"/dkv/kpd?tab={emp}&quarter={row['quarter']}&error=Нет+прав")
    
    # Если новый файл загружен, используем его; иначе сохраняем старый
    fdata = await file.read() if file and file.filename else row["file_data"]
    fname = file.filename if file and file.filename else row["file_name"]
    fmime = file.content_type if file and file.filename else row["file_mime"]
    
    db.execute("""UPDATE dkv_kpd SET 
        name=?, unit=?, formula=?, formula_desc=?, data_source=?, weight=?, 
        value_type=?, threshold=?, goal=?, challenge=?, fact=?,
        file_name=?, file_data=?, file_mime=?
        WHERE id=?""",
        (name, unit, formula, formula_desc, data_source, weight, value_type,
         threshold, goal, challenge, fact, fname, fdata, fmime, kid))
    db.commit()
    employee = row["employee"]
    qtr = row["quarter"]
    db.close()
    return redir(f"/dkv/kpd?tab={employee}&quarter={qtr}")

@app.post("/dkv/kpd/file_delete/{kid}")
async def dkv_kpd_file_delete(kid: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    row = db.execute("SELECT * FROM dkv_kpd WHERE id=?", (kid,)).fetchone()
    if not row: raise HTTPException(404)
    db.execute("UPDATE dkv_kpd SET file_name=NULL, file_data=NULL, file_mime=NULL WHERE id=?", (kid,))
    db.commit()
    employee = row["employee"]
    qtr = row["quarter"]
    db.close()
    return redir(f"/dkv/kpd?tab={employee}&quarter={qtr}")

@app.get("/dkv/kpd/file/{kid}")
async def dkv_kpd_file(kid: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    row = db.execute("SELECT * FROM dkv_kpd WHERE id=?", (kid,)).fetchone()
    db.close()
    if not row or not row["file_data"]: raise HTTPException(404)
    from urllib.parse import quote
    fname_encoded = quote(row["file_name"] or "document")
    return StreamingResponse(io.BytesIO(row["file_data"]),
        media_type=row["file_mime"] or "application/octet-stream",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{fname_encoded}"})

# ── Import ───────────────────────────────────────────────────
@app.get("/org/{org_id}/import/template/{itype}")
async def download_import_template(org_id: int, itype: str, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = Workbook()
    ws = wb.active
    templates = {
        "sessions": {
            "headers": ["Дата (ДД.ММ.ГГГГ)", "Форма (Очное/Заочное)", "Вид (Очередное/Внеочередное)", "Вопрос 1", "Статус 1 (Рассмотрен/Не рассмотрен)", "Вопрос 2", "Статус 2"],
            "example": ["23.01.2026", "Очное", "Очередное", "Об утверждении отчёта за 2025 год", "Рассмотрен", "О назначении члена Правления", "Рассмотрен"]
        },
        "members": {
            "headers": ["ФИО", "Роль (Председатель СД/Член СД)", "Должность", "Независимый (Да/Нет)", "Дата с (ДД.ММ.ГГГГ)", "Дата по (ДД.ММ.ГГГГ)", "Реквизиты акта"],
            "example": ["Иванов Иван Иванович", "Председатель СД", "Независимый директор", "Да", "01.01.2024", "01.01.2027", "Решение ЕА № 1/24"]
        },
        "board": {
            "headers": ["ФИО", "Должность", "Дата с (ДД.ММ.ГГГГ)", "Дата по (ДД.ММ.ГГГГ)", "Реквизиты акта"],
            "example": ["Петров Пётр Петрович", "Председатель Правления", "01.03.2023", "01.03.2026", "Решение ЕА № 5/23"]
        },
        "accountable": {
            "headers": ["ФИО", "Роль (Руководитель/Работник)", "Служба/Организация", "Должность", "Телефон", "Email", "Дата с", "Дата по"],
            "example": ["Сидоров Сидор Сидорович", "Руководитель", "Служба внутреннего аудита", "Руководитель СВА", "+7 701 123 45 67", "sidorov@example.kz", "01.01.2024", "01.01.2026"]
        },
        "ea": {
            "headers": ["Наименование решения", "Дата решения (ДД.ММ.ГГГГ)", "Наименование вопроса", "Реквизиты протокола"],
            "example": ["Решение ЕА № 12/25 от 15.12.2025", "15.12.2025", "Об утверждении Положения о СД", "Протокол № 56/25 от 03.12.2025"]
        },
        "committees": {
            "headers": ["Название комитета", "ФИО", "Роль (Председатель комитета/Член комитета)"],
            "example": ["Комитет по аудиту и рискам", "Иванов Иван Иванович", "Председатель комитета"]
        },
    }
    tmpl = templates.get(itype)
    if not tmpl: raise HTTPException(404)
    from openpyxl.styles import Font, PatternFill, Alignment
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2D3A1F")
    example_fill = PatternFill("solid", fgColor="F0F7E8")
    for col, h in enumerate(tmpl["headers"], 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font; cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        ws.column_dimensions[cell.column_letter].width = max(18, len(h)//2+5)
    for col, v in enumerate(tmpl["example"], 1):
        cell = ws.cell(row=2, column=col, value=v)
        cell.fill = example_fill; cell.font = Font(italic=True, color="555555")
    ws.row_dimensions[1].height = 36
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="template_{itype}.xlsx"'})

@app.post("/org/{org_id}/import/{itype}")
async def do_import(org_id: int, itype: str, request: Request, file: UploadFile = File(None)):
    err, s = guard(request, org_id)
    if err: return err
    if not file or not file.filename: return redir(f"/org/{org_id}?tab=import")
    import openpyxl
    data = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(data))
    ws = wb.active
    rows = [r for r in ws.iter_rows(min_row=2, values_only=True) if any(r)]
    db = get_db(); count = 0
    for row in rows:
        try:
            if itype == "members" and row[0]:
                is_ind = 1 if str(row[3] or "").strip().lower() in ["да","yes","1"] else 0
                db.execute("INSERT INTO sd_members (org_id,full_name,role,is_independent,position,date_from,date_to,decision,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
                    (org_id, str(row[0]).strip(), str(row[1] or "mem").strip(), is_ind, str(row[2] or "").strip(), str(row[4] or "").strip(), str(row[5] or "").strip(), str(row[6] or "").strip(), now(), now())); count+=1
            elif itype == "board" and row[0]:
                db.execute("INSERT INTO board_members (org_id,full_name,position,date_from,date_to,decision,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?)",
                    (org_id, str(row[0]).strip(), str(row[1] or "").strip(), str(row[2] or "").strip(), str(row[3] or "").strip(), str(row[4] or "").strip(), now(), now())); count+=1
            elif itype == "ea" and row[0]:
                db.execute("INSERT INTO ea_decisions (org_id,title,decision_date,question_text,requisites,created_at,updated_at) VALUES (?,?,?,?,?,?,?)",
                    (org_id, str(row[0]).strip(), str(row[1] or "").strip(), str(row[2] or "").strip(), str(row[3] or "").strip(), now(), now())); count+=1
            elif itype == "sessions" and row[0]:
                res = db.execute("INSERT INTO sd_sessions (org_id,session_date,format,order_type,created_at) VALUES (?,?,?,?,?)",
                    (org_id, str(row[0]).strip(), str(row[1] or "Очное").strip(), str(row[2] or "Очередное").strip(), now()))
                sid = res.lastrowid; item_order = 1
                for qi in range(3, len(row)-1, 2):
                    if row[qi]:
                        status = "done" if str(row[qi+1] or "").lower() in ["рассмотрен","done"] else "notdone"
                        db.execute("INSERT INTO sd_agenda_items (session_id,item_order,text,status,note) VALUES (?,?,?,?,?)", (sid, item_order, str(row[qi]).strip(), status, "")); item_order+=1
                count+=1
            elif itype == "committees" and row[0] and row[1]:
                cmt_name = str(row[0]).strip()
                member_name = str(row[1]).strip()
                member_role = str(row[2] or "Член комитета").strip()
                # Найти или создать комитет
                existing = db.execute("SELECT id FROM committees WHERE org_id=? AND name=?", (org_id, cmt_name)).fetchone()
                if existing:
                    cmt_id = existing["id"]
                else:
                    res = db.execute("INSERT INTO committees (org_id,name,color,created_at) VALUES (?,?,?,?)",
                        (org_id, cmt_name, "#1a6b3c", now()))
                    cmt_id = res.lastrowid
                db.execute("INSERT INTO committee_members (committee_id,full_name,role) VALUES (?,?,?)",
                    (cmt_id, member_name, member_role))
                count+=1
            elif itype == "accountable" and row[0]:
                pos_raw = str(row[1] or "").strip().lower()
                pos = "head" if pos_raw in ["руководитель", "head"] else "emp"
                db.execute("INSERT INTO accountable (org_id,full_name,position,org_name,phone,email,date_from,date_to,decision,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                    (org_id, str(row[0]).strip(), pos, str(row[2] or "").strip(), str(row[4] or "").strip(),
                     str(row[5] or "").strip(), str(row[6] or "").strip(), str(row[7] or "").strip(),
                     "", now(), now()))
                count+=1
        except: pass
    db.commit(); db.close()
    log_activity(s, f"import_{itype}", org_id, f"Импортировано {count} записей")
    return redir(f"/org/{org_id}?tab=import")

# ── Excel Export ───────────────────────────────────────────
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import io

def xl_style(wb_sheet, row, col, value, bold=False, bg=None, color="000000", center=False, wrap=False, size=11):
    cell = wb_sheet.cell(row=row, column=col, value=value)
    cell.font = Font(name="Arial", bold=bold, color=color, size=size)
    if bg:
        cell.fill = PatternFill("solid", start_color=bg)
    if center:
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=wrap)
    else:
        cell.alignment = Alignment(vertical="center", wrap_text=True)
    thin = Side(style="thin", color="CCCCCC")
    cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    return cell

def make_excel_response(wb, filename):
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    from fastapi.responses import Response
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="export.xlsx"'}
    )

# ── Export: Состав СД ──────────────────────────────────────
@app.get("/org/{org_id}/export/members")
async def export_members(org_id: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    org = ORG_MAP.get(org_id)
    if not org: raise HTTPException(404)
    db = get_db()
    members = db.execute("SELECT * FROM sd_members WHERE org_id=? ORDER BY id", (org_id,)).fetchall()
    db.close()

    wb = Workbook()
    ws = wb.active
    ws.title = "Состав СД"

    # Title
    ws.merge_cells("A1:G1")
    xl_style(ws, 1, 1, f"{org['name']} — Состав Совета директоров", bold=True, bg="2D3A1F", color="F5F0E8", center=True, size=12)
    ws.row_dimensions[1].height = 22

    # Subtitle
    ws.merge_cells("A2:G2")
    from datetime import date
    xl_style(ws, 2, 1, f"Выгружено: {date.today().strftime('%d.%m.%Y')} · ДКВ АО «НИХ» Байтерек", bg="F2F2F2", color="888888", center=True, size=9)

    # Headers
    headers = ["№", "ФИО", "Роль", "Должность", "Дата начала", "Дата окончания", "Реквизиты решения"]
    for ci, h in enumerate(headers, 1):
        xl_style(ws, 3, ci, h, bold=True, bg="1F4E2C", color="FFFFFF", center=True)
    ws.row_dimensions[3].height = 18

    # Data
    for ri, m in enumerate(members, 1):
        row = 3 + ri
        bg = "F2F8EC" if ri % 2 == 0 else "FFFFFF"
        role_text = "Председатель СД" if m["role"] == "psd" else "Член СД"
        if m["is_independent"]:
            role_text += " (Независимый директор)"
        xl_style(ws, row, 1, ri, center=True, bg=bg)
        xl_style(ws, row, 2, m["full_name"], bold=(m["role"]=="psd"), bg=bg)
        xl_style(ws, row, 3, role_text, bg=bg)
        xl_style(ws, row, 4, m["position"] or "—", bg=bg)
        xl_style(ws, row, 5, fmt(m["date_from"]) if m["date_from"] else "—", center=True, bg=bg)
        xl_style(ws, row, 6, fmt(m["date_to"]) if m["date_to"] else "н.в.", center=True, bg=bg)
        xl_style(ws, row, 7, m["decision"] or "—", bg=bg)
        ws.row_dimensions[row].height = 16

    # Summary
    total = len(members)
    ind = sum(1 for m in members if m["is_independent"])
    summary_row = 4 + total
    ws.merge_cells(f"A{summary_row}:G{summary_row}")
    xl_style(ws, summary_row, 1, f"Итого: {total} членов · из них независимых директоров: {ind}", bg="C6E0B4", color="276221", bold=True)

    # Column widths
    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 28
    ws.column_dimensions["D"].width = 40
    ws.column_dimensions["E"].width = 14
    ws.column_dimensions["F"].width = 14
    ws.column_dimensions["G"].width = 30

    short = org["short"].replace(" ", "_")
    return make_excel_response(wb, f"{short}_Состав_СД.xlsx")

# ── Export: Заседания СД ───────────────────────────────────
@app.get("/org/{org_id}/export/sessions")
async def export_sessions(org_id: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    org = ORG_MAP.get(org_id)
    if not org: raise HTTPException(404)
    db = get_db()
    sessions = db.execute("SELECT * FROM sd_sessions WHERE org_id=? ORDER BY session_date DESC", (org_id,)).fetchall()
    agenda_all = db.execute("SELECT * FROM sd_agenda_items ORDER BY session_id, item_order").fetchall()
    db.close()

    agenda_map = {}
    for a in agenda_all:
        agenda_map.setdefault(a["session_id"], []).append(a)

    wb = Workbook()
    ws = wb.active
    ws.title = "Заседания СД"

    ws.merge_cells("A1:F1")
    xl_style(ws, 1, 1, f"{org['name']} — Заседания Совета директоров", bold=True, bg="2D3A1F", color="F5F0E8", center=True, size=12)
    ws.row_dimensions[1].height = 22
    ws.merge_cells("A2:F2")
    from datetime import date
    xl_style(ws, 2, 1, f"Выгружено: {date.today().strftime('%d.%m.%Y')} · ДКВ АО «НИХ» Байтерек", bg="F2F2F2", color="888888", center=True, size=9)

    headers = ["№ засед.", "Дата", "Форма", "Вид", "Вопрос повестки", "Статус"]
    for ci, h in enumerate(headers, 1):
        xl_style(ws, 3, ci, h, bold=True, bg="1F4E2C", color="FFFFFF", center=True)
    ws.row_dimensions[3].height = 18

    row = 4
    total_q = 0
    done_q = 0
    SESS_COLORS = ["E2EFDA", "EAF2FF", "FFF2CC", "FCE4D6", "EDEDED"]
    for si, sess in enumerate(sessions):
        items = agenda_map.get(sess["id"], [])
        sess_bg = SESS_COLORS[si % len(SESS_COLORS)]
        if not items:
            xl_style(ws, row, 1, si+1, center=True, bold=True, bg=sess_bg)
            xl_style(ws, row, 2, fmt(sess["session_date"]), bold=True, bg=sess_bg)
            xl_style(ws, row, 3, sess["format"], center=True, bg=sess_bg)
            xl_style(ws, row, 4, sess["order_type"], center=True, bg=sess_bg)
            xl_style(ws, row, 5, "—", bg=sess_bg)
            xl_style(ws, row, 6, "—", center=True, bg=sess_bg)
            row += 1
        else:
            for qi, a in enumerate(items):
                total_q += 1
                status_text = "✓ Рассмотрен" if a["status"] == "done" else "✗ Не рассмотрен"
                status_bg = "C6EFCE" if a["status"] == "done" else "FFC7CE"
                status_color = "276221" if a["status"] == "done" else "9C0006"
                if done_q is not None and a["status"] == "done":
                    done_q += 1
                xl_style(ws, row, 1, si+1 if qi==0 else "", center=True, bold=True, bg=sess_bg)
                xl_style(ws, row, 2, fmt(sess["session_date"]) if qi==0 else "", bold=True, bg=sess_bg)
                xl_style(ws, row, 3, sess["format"] if qi==0 else "", center=True, bg=sess_bg)
                xl_style(ws, row, 4, sess["order_type"] if qi==0 else "", center=True, bg=sess_bg)
                xl_style(ws, row, 5, a["text"] or "—", bg=sess_bg)
                xl_style(ws, row, 6, status_text, center=True, bg=status_bg, color=status_color, bold=True)
                ws.row_dimensions[row].height = 16
                row += 1

    ws.merge_cells(f"A{row}:F{row}")
    xl_style(ws, row, 1, f"Итого: {len(sessions)} заседаний · {total_q} вопросов · рассмотрено: {done_q}", bg="C6E0B4", color="276221", bold=True)

    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 12
    ws.column_dimensions["D"].width = 16
    ws.column_dimensions["E"].width = 55
    ws.column_dimensions["F"].width = 18

    short = org["short"].replace(" ", "_")
    return make_excel_response(wb, f"{short}_Заседания_СД.xlsx")

# ── Export: Заседания комитетов ────────────────────────────
@app.get("/org/{org_id}/export/committees")
async def export_committees(org_id: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    org = ORG_MAP.get(org_id)
    if not org: raise HTTPException(404)
    db = get_db()
    cmt_sessions = db.execute("SELECT * FROM committee_sessions WHERE org_id=? ORDER BY session_date ASC", (org_id,)).fetchall()
    participants = db.execute("SELECT * FROM committee_session_participants").fetchall()
    agenda_items = db.execute("SELECT * FROM committee_agenda_items ORDER BY committee_session_id, committee_id, item_order").fetchall()
    cmts = db.execute("SELECT * FROM committees WHERE org_id=?", (org_id,)).fetchall()
    db.close()

    cmt_dict = {c["id"]: c for c in cmts}
    sess_cmt_map = {}
    for p in participants:
        sess_cmt_map.setdefault(p["committee_session_id"], []).append(p["committee_id"])
    sess_agenda_map = {}
    for a in agenda_items:
        key = (a["committee_session_id"], a["committee_id"])
        sess_agenda_map.setdefault(key, []).append(a)

    wb = Workbook()
    ws = wb.active
    ws.title = "Заседания комитетов"

    ws.merge_cells("A1:G1")
    xl_style(ws, 1, 1, f"{org['name']} — Заседания комитетов при СД", bold=True, bg="2D3A1F", color="F5F0E8", center=True, size=12)
    ws.row_dimensions[1].height = 22
    ws.merge_cells("A2:G2")
    from datetime import date
    xl_style(ws, 2, 1, f"Выгружено: {date.today().strftime('%d.%m.%Y')} · ДКВ АО «НИХ» Байтерек", bg="F2F2F2", color="888888", center=True, size=9)

    headers = ["№ засед.", "Дата", "№ прот.", "Комитет", "№ вопроса", "Вопрос повестки", "Статус"]
    for ci, h in enumerate(headers, 1):
        xl_style(ws, 3, ci, h, bold=True, bg="1F4E2C", color="FFFFFF", center=True)
    ws.row_dimensions[3].height = 18

    CMT_COLORS = ["E2EFDA","EAF2FF","FFF2CC","FCE4D6","E8E0F0","FDEBD0","D5F5E3"]
    row = 4
    total_sess = len(cmt_sessions)
    for si, sess in enumerate(cmt_sessions):
        cmts_in_sess = sess_cmt_map.get(sess["id"], [])
        sess_first = True
        for cid in cmts_in_sess:
            c = cmt_dict.get(cid)
            if not c: continue
            cmt_bg = CMT_COLORS[list(cmt_dict.keys()).index(cid) % len(CMT_COLORS)] if cid in cmt_dict else "EEEEEE"
            agenda = sess_agenda_map.get((sess["id"], cid), [])
            if not agenda:
                xl_style(ws, row, 1, si+1 if sess_first else "", center=True, bold=True)
                xl_style(ws, row, 2, fmt(sess["session_date"]) if sess_first else "", bold=True)
                xl_style(ws, row, 3, sess["protocol_num"] or "—" if sess_first else "", center=True)
                xl_style(ws, row, 4, c["name"], bold=True, bg=cmt_bg)
                xl_style(ws, row, 5, "—", center=True, bg=cmt_bg)
                xl_style(ws, row, 6, "—", bg=cmt_bg)
                xl_style(ws, row, 7, "—", center=True, bg=cmt_bg)
                row += 1
                sess_first = False
            else:
                for qi, a in enumerate(agenda):
                    status_text = "✓ Рассмотрен" if a["status"] == "done" else "✗ Не рассмотрен"
                    status_bg = "C6EFCE" if a["status"] == "done" else "FFC7CE"
                    status_color = "276221" if a["status"] == "done" else "9C0006"
                    xl_style(ws, row, 1, si+1 if sess_first and qi==0 else "", center=True, bold=True)
                    xl_style(ws, row, 2, fmt(sess["session_date"]) if sess_first and qi==0 else "", bold=True)
                    xl_style(ws, row, 3, sess["protocol_num"] or "—" if sess_first and qi==0 else "", center=True)
                    xl_style(ws, row, 4, c["name"] if qi==0 else "", bold=True, bg=cmt_bg)
                    xl_style(ws, row, 5, qi+1, center=True, bg=cmt_bg)
                    xl_style(ws, row, 6, a["text"] or "—", bg=cmt_bg)
                    xl_style(ws, row, 7, status_text, center=True, bg=status_bg, color=status_color, bold=True)
                    ws.row_dimensions[row].height = 16
                    row += 1
                    if qi == 0:
                        sess_first = False

    ws.merge_cells(f"A{row}:G{row}")
    xl_style(ws, row, 1, f"Итого: {total_sess} заседаний · {len(cmt_dict)} комитетов", bg="C6E0B4", color="276221", bold=True)

    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 10
    ws.column_dimensions["D"].width = 25
    ws.column_dimensions["E"].width = 10
    ws.column_dimensions["F"].width = 50
    ws.column_dimensions["G"].width = 18

    short = org["short"].replace(" ", "_")
    return make_excel_response(wb, f"{short}_Заседания_комитетов.xlsx")

# ── Activity logger ────────────────────────────────────────
def log_activity(session: dict, action: str, org_id: int = None, details: str = ""):
    try:
        db = get_db()
        db.execute(
            "INSERT INTO activity_log (user_id,username,full_name,action,org_id,details,created_at) VALUES (?,?,?,?,?,?,?)",
            (session["uid"], session["username"], session["full_name"], action, org_id, details, now()))
        db.execute(
            "INSERT OR REPLACE INTO user_sessions (user_id,username,full_name,last_seen) VALUES (?,?,?,?)",
            (session["uid"], session["username"], session["full_name"], now()))
        db.commit(); db.close()
    except: pass

# ── Update session last_seen on each request ───────────────
from starlette.middleware.base import BaseHTTPMiddleware
class SessionTracker(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        s = get_session(request)
        if s:
            try:
                db = get_db()
                db.execute("INSERT OR REPLACE INTO user_sessions (user_id,username,full_name,last_seen) VALUES (?,?,?,?)",
                    (s["uid"], s["username"], s["full_name"], now()))
                db.commit(); db.close()
            except: pass
        return await call_next(request)
app.add_middleware(SessionTracker)

# ── Static files ───────────────────────────────────────────
from fastapi.staticfiles import StaticFiles
import os
if os.path.exists("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")

# ── Plan работы СД ─────────────────────────────────────────
@app.post("/org/{org_id}/workplan/upload")
async def upload_workplan(org_id: int, request: Request, file: UploadFile = File(...)):
    err, s = guard(request, org_id)
    if err: return err
    fdata = await file.read()
    db = get_db()
    db.execute("""INSERT OR REPLACE INTO sd_work_plan (org_id,file_name,file_data,file_mime,uploaded_at,updated_at)
        VALUES (?,?,?,?,?,?)""", (org_id, file.filename, fdata, file.content_type, now(), now()))
    db.commit(); db.close()
    log_activity(s, "upload_workplan", org_id, file.filename)
    return redir(f"/org/{org_id}?tab=workplan")

@app.post("/org/{org_id}/workplan/delete")
async def delete_workplan(org_id: int, request: Request):
    err, s = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("DELETE FROM sd_work_plan WHERE org_id=?", (org_id,))
    db.commit(); db.close()
    log_activity(s, "delete_workplan", org_id)
    return redir(f"/org/{org_id}?tab=workplan")

@app.get("/org/{org_id}/workplan/view")
async def view_workplan(org_id: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    row = db.execute("SELECT * FROM sd_work_plan WHERE org_id=?", (org_id,)).fetchone()
    db.close()
    if not row or not row["file_data"]: raise HTTPException(404)
    return StreamingResponse(io.BytesIO(row["file_data"]),
        media_type=row["file_mime"] or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="workplan.pdf"'})

# ── Решения ЕА ─────────────────────────────────────────────
@app.post("/org/{org_id}/ea/add")
async def add_ea(org_id: int, request: Request,
    title: str = Form(...), decision_date: str = Form(""),
    question_text: str = Form(""), requisites: str = Form(""),
    file: UploadFile = File(None)):
    err, s = guard(request, org_id)
    if err: return err
    fname = fdata = fmime = None
    if file and file.filename:
        fname = file.filename; fdata = await file.read(); fmime = file.content_type
    db = get_db()
    db.execute("INSERT INTO ea_decisions (org_id,title,decision_date,question_text,requisites,file_name,file_data,file_mime,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
        (org_id, title, decision_date, question_text, requisites, fname, fdata, fmime, now(), now()))
    db.commit(); db.close()
    log_activity(s, "add_ea_decision", org_id, title)
    return redir(f"/org/{org_id}?tab=ea")

@app.post("/org/{org_id}/ea/edit/{eid}")
async def edit_ea(org_id: int, eid: int, request: Request,
    title: str = Form(...), decision_date: str = Form(""),
    question_text: str = Form(""), requisites: str = Form(""),
    delete_file: str = Form(""), file: UploadFile = File(None)):
    err, s = guard(request, org_id)
    if err: return err
    db = get_db()
    # Handle file
    if file and file.filename:
        fdata = await file.read()
        db.execute("UPDATE ea_decisions SET title=?,decision_date=?,question_text=?,requisites=?,file_name=?,file_data=?,file_mime=?,updated_at=? WHERE id=? AND org_id=?",
            (title, decision_date, question_text, requisites, file.filename, fdata, file.content_type, now(), eid, org_id))
    elif delete_file == "1":
        db.execute("UPDATE ea_decisions SET title=?,decision_date=?,question_text=?,requisites=?,file_name=NULL,file_data=NULL,file_mime=NULL,updated_at=? WHERE id=? AND org_id=?",
            (title, decision_date, question_text, requisites, now(), eid, org_id))
    else:
        db.execute("UPDATE ea_decisions SET title=?,decision_date=?,question_text=?,requisites=?,updated_at=? WHERE id=? AND org_id=?",
            (title, decision_date, question_text, requisites, now(), eid, org_id))
    db.commit(); db.close()
    log_activity(s, "edit_ea_decision", org_id, title)
    return redir(f"/org/{org_id}?tab=ea")

@app.post("/org/{org_id}/ea/delete/{eid}")
async def delete_ea(org_id: int, eid: int, request: Request):
    err, s = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("DELETE FROM ea_decisions WHERE id=? AND org_id=?", (eid, org_id))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=ea")

@app.get("/org/{org_id}/ea/view/{eid}")
async def view_ea(org_id: int, eid: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    row = db.execute("SELECT * FROM ea_decisions WHERE id=? AND org_id=?", (eid, org_id)).fetchone()
    db.close()
    if not row or not row["file_data"]: raise HTTPException(404)
    return StreamingResponse(io.BytesIO(row["file_data"]),
        media_type=row["file_mime"] or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="ea_decision.pdf"'})

@app.get("/org/{org_id}/ea/letter/{eid}", response_class=HTMLResponse)
async def ea_letter(org_id: int, eid: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    org = ORG_MAP[org_id]
    db = get_db()
    ea = db.execute("SELECT * FROM ea_decisions WHERE id=? AND org_id=?", (eid, org_id)).fetchone()
    db.close()
    if not ea: raise HTTPException(404)
    return templates.TemplateResponse("ea_letter.html", {
        "request": request, "session": s, "org": org, "ea": ea
    })

# ── Файлы к заседаниям СД ─────────────────────────────────
@app.post("/org/{org_id}/sessions/{sid}/file/upload")
async def upload_sess_file(org_id: int, sid: int, request: Request, file: UploadFile = File(...)):
    err, s = guard(request, org_id)
    if err: return err
    fdata = await file.read()
    db = get_db()
    db.execute("INSERT INTO sd_session_files (session_id,file_name,file_data,file_mime,uploaded_at) VALUES (?,?,?,?,?)",
        (sid, file.filename, fdata, file.content_type, now()))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=sessions")

@app.post("/org/{org_id}/sessions/{sid}/file/delete/{fid}")
async def delete_sess_file(org_id: int, sid: int, fid: int, request: Request):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("DELETE FROM sd_session_files WHERE id=? AND session_id=?", (fid, sid))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=sessions")

@app.get("/org/{org_id}/sessions/{sid}/file/view/{fid}")
async def view_sess_file(org_id: int, sid: int, fid: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    row = db.execute("SELECT * FROM sd_session_files WHERE id=? AND session_id=?", (fid, sid)).fetchone()
    db.close()
    if not row or not row["file_data"]: raise HTTPException(404)
    return StreamingResponse(io.BytesIO(row["file_data"]),
        media_type=row["file_mime"] or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="session_doc.pdf"'})

# ── Файлы к заседаниям комитетов ──────────────────────────
@app.post("/org/{org_id}/cmt-sessions/{sid}/file/upload")
async def upload_cmt_file(org_id: int, sid: int, request: Request, file: UploadFile = File(...)):
    err, _ = guard(request, org_id)
    if err: return err
    fdata = await file.read()
    db = get_db()
    db.execute("INSERT INTO cmt_session_files (committee_session_id,file_name,file_data,file_mime,uploaded_at) VALUES (?,?,?,?,?)",
        (sid, file.filename, fdata, file.content_type, now()))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees&ctab=sessions")

@app.post("/org/{org_id}/cmt-sessions/{sid}/file/delete/{fid}")
async def delete_cmt_file(org_id: int, sid: int, fid: int, request: Request):
    err, _ = guard(request, org_id)
    if err: return err
    db = get_db()
    db.execute("DELETE FROM cmt_session_files WHERE id=? AND committee_session_id=?", (fid, sid))
    db.commit(); db.close()
    return redir(f"/org/{org_id}?tab=committees&ctab=sessions")

@app.get("/org/{org_id}/cmt-sessions/{sid}/file/view/{fid}")
async def view_cmt_file(org_id: int, sid: int, fid: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    db = get_db()
    row = db.execute("SELECT * FROM cmt_session_files WHERE id=? AND committee_session_id=?", (fid, sid)).fetchone()
    db.close()
    if not row or not row["file_data"]: raise HTTPException(404)
    return StreamingResponse(io.BytesIO(row["file_data"]),
        media_type=row["file_mime"] or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="cmt_doc.pdf"'})

# ── Admin: онлайн пользователи и лог ──────────────────────
@app.get("/admin/activity", response_class=HTMLResponse)
async def activity_page(request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    # Curators can also see activity for their orgs
    if s["role"] == "viewer": return redir("/")
    db = get_db()
    online = db.execute("SELECT * FROM user_sessions ORDER BY last_seen DESC").fetchall()
    if s["role"] == "admin":
        logs = db.execute("SELECT * FROM activity_log ORDER BY created_at DESC LIMIT 100").fetchall()
    else:
        # Show only logs for curator's orgs
        curator_org_ids = [o["id"] for o in ORGANIZATIONS if o["curator"].lower() == s["username"].lower()]
        if curator_org_ids:
            placeholders = ",".join("?" * len(curator_org_ids))
            logs = db.execute(f"SELECT * FROM activity_log WHERE org_id IN ({placeholders}) ORDER BY created_at DESC LIMIT 100", curator_org_ids).fetchall()
        else:
            logs = []
    db.close()
    return templates.TemplateResponse("activity.html", {
        "request": request, "session": s,
        "online": online, "logs": logs
    })

# ── Export: Состав комитетов ───────────────────────────────
@app.get("/org/{org_id}/export/committee-members")
async def export_cmt_members(org_id: int, request: Request):
    s = get_session(request)
    if not s: return redir("/login")
    org = ORG_MAP.get(org_id)
    if not org: raise HTTPException(404)
    db = get_db()
    cmts = db.execute("SELECT * FROM committees WHERE org_id=? ORDER BY id", (org_id,)).fetchall()
    members = db.execute("""SELECT cm.*, c.name as cmt_name FROM committee_members cm
        JOIN committees c ON c.id=cm.committee_id
        WHERE c.org_id=? ORDER BY cm.committee_id, cm.id""", (org_id,)).fetchall()
    db.close()
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    wb = Workbook()
    ws = wb.active
    ws.title = "Состав комитетов"
    ws.merge_cells("A1:D1")
    c = ws.cell(1,1,f"{org['name']} — Состав комитетов при СД")
    c.font = Font(name="Arial",bold=True,color="F5F0E8",size=12)
    c.fill = PatternFill("solid",start_color="2D3A1F")
    c.alignment = Alignment(horizontal="center",vertical="center")
    ws.row_dimensions[1].height = 22
    from datetime import date
    ws.merge_cells("A2:D2")
    c2 = ws.cell(2,1,f"Выгружено: {date.today().strftime('%d.%m.%Y')} · ДКВ АО «НИХ» Байтерек")
    c2.font = Font(name="Arial",size=9,color="888888")
    c2.fill = PatternFill("solid",start_color="F2F2F2")
    c2.alignment = Alignment(horizontal="center")
    thin = Side(style="thin",color="CCCCCC")
    brd = Border(left=thin,right=thin,top=thin,bottom=thin)
    hdrs = ["№","Комитет","ФИО","Роль"]
    for ci,h in enumerate(hdrs,1):
        cell = ws.cell(3,ci,h)
        cell.font = Font(name="Arial",bold=True,color="FFFFFF",size=10)
        cell.fill = PatternFill("solid",start_color="1F4E2C")
        cell.alignment = Alignment(horizontal="center",vertical="center")
        cell.border = brd
    row = 4
    CMT_COLORS = ["E2EFDA","EAF2FF","FFF2CC","FCE4D6","E8E0F0"]
    cmt_idx = {c["id"]:i for i,c in enumerate(cmts)}
    for ri,m in enumerate(members,1):
        bg = CMT_COLORS[cmt_idx.get(m["committee_id"],0) % len(CMT_COLORS)]
        for ci,val in enumerate([ri, m["cmt_name"], m["full_name"], m["role"]],1):
            cell = ws.cell(row,ci,val)
            cell.font = Font(name="Arial",size=10,bold=(ci==2))
            cell.fill = PatternFill("solid",start_color=bg)
            cell.alignment = Alignment(vertical="center",wrap_text=True)
            cell.border = brd
        ws.row_dimensions[row].height=16
        row+=1
    ws.merge_cells(f"A{row}:D{row}")
    s_cell = ws.cell(row,1,f"Итого: {len(cmts)} комитетов · {len(members)} членов")
    s_cell.font = Font(name="Arial",bold=True,color="276221",size=10)
    s_cell.fill = PatternFill("solid",start_color="C6E0B4")
    ws.column_dimensions["A"].width=5
    ws.column_dimensions["B"].width=25
    ws.column_dimensions["C"].width=35
    ws.column_dimensions["D"].width=22
    from fastapi.responses import Response
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return Response(content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="export.xlsx"'})

# ── Forgot Password ────────────────────────────────────────
@app.get("/forgot-password", response_class=HTMLResponse)
async def forgot_page(request: Request, sent: str = ""):
    return templates.TemplateResponse("forgot.html", {"request": request, "sent": sent})

@app.post("/forgot-password")
async def forgot_submit(request: Request,
    username: str = Form(""), comment: str = Form("")):
    db = get_db()
    user = db.execute("SELECT * FROM users WHERE LOWER(username)=LOWER(?)", (username,)).fetchone()
    if user:
        db.execute("INSERT INTO password_requests (username, full_name, comment, status, created_at) VALUES (?,?,?,?,?)",
            (user["username"], user["full_name"], comment, "new", now()))
        db.commit()
    db.close()
    return redir("/forgot-password?sent=1")

@app.post("/users/pwd-request/resolve/{rid}")
async def resolve_pwd_request(rid: int, request: Request,
    new_password: str = Form(...)):
    s = get_session(request)
    if not s or s["role"] != "admin": return redir("/login")
    import bcrypt
    db = get_db()
    req = db.execute("SELECT * FROM password_requests WHERE id=?", (rid,)).fetchone()
    if req:
        pw_hash = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode()
        db.execute("UPDATE users SET password_hash=? WHERE username=?", (pw_hash, req["username"]))
        db.execute("UPDATE password_requests SET status='done' WHERE id=?", (rid,))
        db.commit()
    db.close()
    return redir("/users")

@app.post("/users/pwd-request/dismiss/{rid}")
async def dismiss_pwd_request(rid: int, request: Request):
    s = get_session(request)
    if not s or s["role"] != "admin": return redir("/login")
    db = get_db()
    db.execute("DELETE FROM password_requests WHERE id=?", (rid,))
    db.commit()
    db.close()
    return redir("/users")

# ── Парсинг повестки СД из .docx ──────────────────────────
@app.post("/org/{org_id}/sessions/parse-agenda")
async def parse_agenda_docx(org_id: int, request: Request, file: UploadFile = File(...)):
    """Парсит .docx повестки дня и возвращает структурированные данные JSON"""
    s = get_session(request)
    if not s: return JSONResponse({"error": "Не авторизован"}, status_code=401)

    import re
    try:
        import docx as _docx
    except ImportError:
        import subprocess
        subprocess.run(["pip", "install", "python-docx", "--break-system-packages", "-q"])
        import docx as _docx

    fdata = await file.read()
    doc = _docx.Document(io.BytesIO(fdata))

    # Извлекаем параграфы
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    # --- Парсинг даты ---
    date_iso = ""
    MONTHS_RU = {
        "января":1,"февраля":2,"марта":3,"апреля":4,"мая":5,"июня":6,
        "июля":7,"августа":8,"сентября":9,"октября":10,"ноября":11,"декабря":12
    }
    # Шаблон: «30» апреля 2026 года  или  30 апреля 2026
    dm = re.search(r'[«"]?(\d{1,2})[»"]?\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+(\d{4})', full_text, re.IGNORECASE)
    if dm:
        day, mon_str, year = int(dm.group(1)), dm.group(2).lower(), int(dm.group(3))
        month = MONTHS_RU.get(mon_str, 0)
        if month:
            date_iso = f"{year}-{month:02d}-{day:02d}"

    # --- Тип заседания ---
    format_val = "Очное"
    if re.search(r'заочн', full_text, re.IGNORECASE): format_val = "Заочное"
    elif re.search(r'смешанн', full_text, re.IGNORECASE): format_val = "Смешанное"

    order_type = "Очередное"
    if re.search(r'внеочередн', full_text, re.IGNORECASE): order_type = "Внеочередное"

    # --- Время и место ---
    time_val = ""
    tm = re.search(r'(\d{1,2})\s*[чh]\s*[:\.\s]*\s*(\d{2})\s*[мm]', full_text, re.IGNORECASE)
    if tm: time_val = f"{int(tm.group(1)):02d}:{tm.group(2)}"

    place_val = ""
    for p in paragraphs:
        if re.search(r'место проведения', p, re.IGNORECASE):
            # Следующий параграф содержит адрес
            idx = paragraphs.index(p)
            if idx + 1 < len(paragraphs):
                place_val = paragraphs[idx + 1]
            break
        # Или адрес сразу в этой строке
        if re.search(r'\d{6}|г\.\s*Астана|г\.\s*Алматы', p) and 'Место' not in p:
            place_val = p
            break

    # --- Вопросы повестки из таблиц ---
    agenda_items = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                text = cells[1] if len(cells) > 1 else cells[0]
                # Пропускаем заголовок
                if text and text not in ("Наименование вопроса", "№", "Вопрос"):
                    agenda_items.append({
                        "text": text,
                        "status": "done",
                        "note": cells[2] if len(cells) > 2 else ""
                    })

    return JSONResponse({
        "date": date_iso,
        "format": format_val,
        "order_type": order_type,
        "time": time_val,
        "place": place_val,
        "agenda": agenda_items,
        "filename": file.filename
    })

# ── Сохранение распарсенной повестки СД ───────────────────
@app.post("/org/{org_id}/sessions/import-agenda")
async def import_agenda(org_id: int, request: Request,
    session_date: str = Form(...),
    format: str = Form("Очное"),
    order_type: str = Form("Очередное"),
    agenda_texts: list = Form([]),
    agenda_statuses: list = Form([]),
    agenda_notes: list = Form([])
):
    err, s = guard(request, org_id)
    if err: return err
    db = get_db()
    cur = db.execute(
        "INSERT INTO sd_sessions (org_id,session_date,format,order_type,created_at,updated_at) VALUES (?,?,?,?,?,?)",
        (org_id, session_date, format, order_type, now(), now())
    )
    sid = cur.lastrowid
    for i, text in enumerate(agenda_texts):
        if text.strip():
            status = agenda_statuses[i] if i < len(agenda_statuses) else "done"
            note = agenda_notes[i] if i < len(agenda_notes) else ""
            db.execute(
                "INSERT INTO sd_agenda_items (session_id,item_order,text,status,note) VALUES (?,?,?,?,?)",
                (sid, i + 1, text.strip(), status, note.strip())
            )
    db.commit(); db.close()
    log_activity(s, "import_agenda_docx", org_id, f"Импорт повестки от {session_date}, {len(agenda_texts)} вопросов")
    return redir(f"/org/{org_id}?tab=sessions")


# ─────────────────────────────────────────────────────────
#  АГЕНТ МОНИТОРИНГА СОСТАВОВ ДО
# ─────────────────────────────────────────────────────────

import difflib as _difflib

def _fuzzy_match(name_a: str, name_b: str) -> bool:
    a = name_a.strip().lower()
    b = name_b.strip().lower()
    if a == b:
        return True
    return _difflib.SequenceMatcher(None, a, b).ratio() >= 0.75


def _compare_members(db_names: list, snap_names: list):
    added = [s for s in snap_names if not any(_fuzzy_match(s, d) for d in db_names)]
    removed = [d for d in db_names if not any(_fuzzy_match(d, s) for s in snap_names)]
    return added, removed


@app.get("/agents", response_class=HTMLResponse)
def agents_page(request: Request):
    s = get_session(request)
    if not s:
        return redir("/login")
    db = get_db()
    org_short_map = {o["id"]: o["short"] for o in ORGANIZATIONS}
    rows = db.execute("""
        SELECT a.id, a.org_id, a.alert_type, a.name, a.is_read, a.created_at,
               sn.source_label, sn.created_by
        FROM agent_alerts a
        JOIN agent_snapshots sn ON a.snapshot_id = sn.id
        ORDER BY a.id DESC LIMIT 50
    """).fetchall()
    alerts = [dict(r, org_short=org_short_map.get(r["org_id"], "?")) for r in rows]
    unread = db.execute("SELECT COUNT(*) FROM agent_alerts WHERE is_read=0").fetchone()[0]
    db.close()
    return templates.TemplateResponse("agents.html", {
        "request": request, "session": s,
        "orgs": ORGANIZATIONS, "alerts": alerts, "unread": unread,
    })


@app.post("/agents/snapshot")
async def agents_snapshot(
    request: Request,
    org_id: int = Form(...),
    source_label: str = Form(""),
    names_raw: str = Form(...),
):
    s = get_session(request)
    if not s:
        return redir("/login")
    lines = [l.strip() for l in names_raw.splitlines() if l.strip()]
    if not lines:
        return redir("/agents")
    db = get_db()
    db_names = [r["full_name"] for r in db.execute(
        "SELECT full_name FROM sd_members WHERE org_id=? AND (date_to IS NULL OR date_to='')",
        (org_id,)
    ).fetchall()]
    added, removed = _compare_members(db_names, lines)
    cur = db.execute(
        "INSERT INTO agent_snapshots (org_id, source_label, names_raw, created_by, created_at) VALUES (?,?,?,?,?)",
        (org_id, source_label or "Ручной ввод", "\n".join(lines), s["username"], now())
    )
    snap_id = cur.lastrowid
    for nm in added:
        db.execute(
            "INSERT INTO agent_alerts (org_id, snapshot_id, alert_type, name, is_read, created_at) VALUES (?,?,?,?,0,?)",
            (org_id, snap_id, "added", nm, now())
        )
    for nm in removed:
        db.execute(
            "INSERT INTO agent_alerts (org_id, snapshot_id, alert_type, name, is_read, created_at) VALUES (?,?,?,?,0,?)",
            (org_id, snap_id, "removed", nm, now())
        )
    db.commit()
    db.close()
    return redir("/agents")


@app.post("/agents/mark-read/{alert_id}")
async def agents_mark_read(request: Request, alert_id: int):
    s = get_session(request)
    if not s:
        return redir("/login")
    db = get_db()
    db.execute("UPDATE agent_alerts SET is_read=1 WHERE id=?", (alert_id,))
    db.commit()
    db.close()
    return redir("/agents")


@app.post("/agents/mark-all-read")
async def agents_mark_all_read(request: Request):
    s = get_session(request)
    if not s:
        return redir("/login")
    db = get_db()
    db.execute("UPDATE agent_alerts SET is_read=1")
    db.commit()
    db.close()
    return redir("/agents")


# ─────────────────────────────────────────────────────────
#  PDF-СПРАВКА ПО ОРГАНИЗАЦИИ
# ─────────────────────────────────────────────────────────

from gen_org_pdf import generate_org_pdf

@app.get("/org/{org_id}/report/pdf")
def org_report_pdf(request: Request, org_id: int):
    s = get_session(request)
    if not s:
        return redir("/login")
    org = ORG_MAP.get(org_id)
    if not org:
        raise HTTPException(404)
    db = get_db()
    try:
        pdf_bytes = generate_org_pdf(org, db)
    finally:
        db.close()
    filename = f"spravka_{org['id']}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
